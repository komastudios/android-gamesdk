#
# Copyright 2019 The Android Open Source Project
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
#
"""
Word summary formatter.
"""

from contextlib import contextmanager
from pathlib import Path
from typing import Optional, TypeVar

from docx import Document
from docx.shared import Inches

from lib.summary_formatters.formatter import SummaryFormatter


class WordFormatter(SummaryFormatter):
    """Word summary formatter."""

    def __init__(self):
        self.__writer = None

    @classmethod
    def is_gdrive_compatible(cls: TypeVar("SummaryFormatter")) -> bool:
        return True

    @contextmanager
    def create_writer(self, summary_path: Path) -> Document:
        self.__writer = Document()

        try:
            yield self.__writer
        finally:
            self.__writer.save(str(summary_path))

    def on_init(self, title: str, summary_utc: str) -> type(None):
        self.__writer.add_heading(title, 0)
        self.__writer.add_paragraph(summary_utc)


        # TODO(baxtermichael@google.com): handle table parsing from markdown
        table = self.__writer.add_table(rows=1, cols=3)
        # table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Frame Timestamp Support'
        hdr_cells[1].text = 'Available'
        hdr_cells[2].text = 'Working'

        records = (
            ('Requested Present Time', 'Y', 'Y'),
            ('Requested Present Time', 'Y', 'Y'),
            ('Rendering Complete Time', 'N', 'n/a'),
            ('Composition Latch Time', 'Y', 'Y'),
            ('First Composition Start Time', 'Y', 'N'),
            ('Last Composition Start Time', 'Y', 'Y'),
            ('First Composition GPU Finished Time', 'Y', 'Y'),
            ('Display Present Time', 'Y', 'Y'),
            ('Dequeue Ready Time', 'Y', 'Y'),
            ('Reads Done Time', 'Y', 'Y'),
        )

        for col1, col2, col3 in records:
            row_cells = table.add_row().cells
            row_cells[0].text = col1
            row_cells[1].text = col2
            row_cells[2].text = col3


    def on_device(self, device_id: str, plot_path: Path,
                  plot_path_relative: Path,
                  summary: Optional[str]) -> type(None):
        self.__writer.add_heading(device_id, 3)
        self.__writer.add_picture(str(plot_path), Inches(6.18))
        if summary is not None:
            self.__writer.add_paragraph(summary)

        self.__writer.add_page_break()

    def on_cross_suite(self, plot_path: Path, plot_path_relative: Path,
                       summary: Optional[str]) -> type(None):
        self.__writer.add_heading("Meta Summary", 2)
        self.__writer.add_picture(str(plot_path), Inches(6.18))
        if summary is not None:
            self.__writer.add_paragraph(summary)

    def on_errors_available(self, title: str) -> type(None):
        self.__writer.add_heading(title, 1)

    def on_device_error(self, device_id: str, summary: str) -> type(None):
        self.__writer.add_heading(device_id, 3)
        self.__writer.add_paragraph(summary)
